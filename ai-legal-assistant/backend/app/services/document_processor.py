"""
Document processing service for extracting text from various file formats
"""

import logging
import os
from typing import Optional
import asyncio
from pathlib import Path

# PDF processing
import pdfplumber
import fitz  # PyMuPDF
import pymupdf4llm

# DOCX processing
from docx import Document as DocxDocument

# OCR processing
import pytesseract
from PIL import Image

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for processing various document formats"""
    
    def __init__(self):
        self.supported_formats = {
            'application/pdf': self._extract_pdf_text,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_docx_text,
            'application/msword': self._extract_doc_text,
            'image/png': self._extract_image_text,
            'image/jpeg': self._extract_image_text,
            'image/tiff': self._extract_image_text
        }
    
    async def extract_text(self, file_path: str, mime_type: str) -> Optional[str]:
        """Extract text from document based on MIME type"""
        try:
            if mime_type not in self.supported_formats:
                logger.warning(f"Unsupported file type: {mime_type}")
                return None
            
            # Run extraction in thread pool to avoid blocking
            loop = asyncio.get_event_loop()
            text = await loop.run_in_executor(
                None, 
                self.supported_formats[mime_type], 
                file_path
            )
            
            if text:
                logger.info(f"Successfully extracted text from {file_path}")
                return text
            else:
                logger.warning(f"No text extracted from {file_path}")
                return None
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return None
    
    def _extract_pdf_text(self, file_path: str) -> Optional[str]:
        """Extract text from PDF using multiple methods"""
        try:
            # Method 1: Try pdfplumber first (better for complex layouts)
            text = self._extract_pdf_with_pdfplumber(file_path)
            if text and len(text.strip()) > 100:
                return text

            # Method 2: Try PyMuPDF (faster, good for simple PDFs)
            text2 = self._extract_pdf_with_pymupdf(file_path)
            if text2 and len(text2.strip()) > 100:
                return text2

            # Method 3: Try pymupdf4llm (specialized for LLM processing)
            text3 = self._extract_pdf_with_pymupdf4llm(file_path)
            if text3 and len(text3.strip()) > 100:
                return text3

            # Fallback: Page-level OCR for pages with little/no text (image-based PDFs)
            ocr_text = self._extract_pdf_with_ocr(file_path)
            if ocr_text and len(ocr_text.strip()) > 50:
                return ocr_text

            logger.warning(f"All PDF extraction methods yielded minimal text for {file_path}")
            return text or text2 or text3 or None
            
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return None
    
    def _extract_pdf_with_pdfplumber(self, file_path: str) -> Optional[str]:
        """Extract text using pdfplumber"""
        try:
            with pdfplumber.open(file_path) as pdf:
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                return '\n'.join(text_parts)
        except Exception as e:
            logger.debug(f"pdfplumber extraction failed: {e}")
            return None
    
    def _extract_pdf_with_pymupdf(self, file_path: str) -> Optional[str]:
        """Extract text using PyMuPDF"""
        try:
            doc = fitz.open(file_path)
            text_parts = []
            for page_num in range(doc.page_count):
                page = doc[page_num]
                text = page.get_text()
                if text:
                    text_parts.append(text)
            doc.close()
            return '\n'.join(text_parts)
        except Exception as e:
            logger.debug(f"PyMuPDF extraction failed: {e}")
            return None

    def _extract_pdf_with_ocr(self, file_path: str) -> Optional[str]:
        """OCR fallback: render pages to images and run Tesseract on low-text PDFs"""
        try:
            doc = fitz.open(file_path)
            text_parts = []
            for page_num in range(doc.page_count):
                page = doc[page_num]
                # Render at higher DPI for better OCR
                pix = page.get_pixmap(dpi=200)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                page_text = pytesseract.image_to_string(img)
                if page_text:
                    text_parts.append(page_text)
            doc.close()
            return '\n'.join(text_parts)
        except Exception as e:
            logger.debug(f"OCR extraction failed: {e}")
            return None
    
    def _extract_pdf_with_pymupdf4llm(self, file_path: str) -> Optional[str]:
        """Extract text using pymupdf4llm"""
        try:
            markdown_text = pymupdf4llm.to_markdown(file_path)
            return markdown_text
        except Exception as e:
            logger.debug(f"pymupdf4llm extraction failed: {e}")
            return None
    
    def _extract_docx_text(self, file_path: str) -> Optional[str]:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            return '\n'.join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return None
    
    def _extract_doc_text(self, file_path: str) -> Optional[str]:
        """Extract text from DOC file (requires additional library)"""
        try:
            # For .doc files, we would need python-docx2txt or similar
            # This is a placeholder implementation
            logger.warning("DOC file extraction not fully implemented")
            return None
        except Exception as e:
            logger.error(f"Error extracting DOC text: {e}")
            return None
    
    def _extract_image_text(self, file_path: str) -> Optional[str]:
        """Extract text from image using OCR"""
        try:
            # Open image
            image = Image.open(file_path)
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Extract text using Tesseract
            text = pytesseract.image_to_string(image)
            
            if text and len(text.strip()) > 10:
                return text
            else:
                logger.warning(f"No text found in image {file_path}")
                return None
                
        except Exception as e:
            logger.error(f"Error extracting image text: {e}")
            return None
    
    def get_document_info(self, file_path: str) -> dict:
        """Get document information without extracting full text"""
        try:
            file_info = {
                "file_size": os.path.getsize(file_path),
                "file_extension": Path(file_path).suffix.lower(),
                "exists": os.path.exists(file_path)
            }
            
            # Try to get page count for PDFs
            if file_path.lower().endswith('.pdf'):
                try:
                    with pdfplumber.open(file_path) as pdf:
                        file_info["page_count"] = len(pdf.pages)
                except:
                    file_info["page_count"] = None
            
            return file_info
            
        except Exception as e:
            logger.error(f"Error getting document info: {e}")
            return {"error": str(e)}
    
    def validate_document(self, file_path: str, mime_type: str) -> dict:
        """Validate document before processing"""
        validation = {
            "is_valid": True,
            "warnings": [],
            "errors": []
        }
        
        try:
            # Check if file exists
            if not os.path.exists(file_path):
                validation["is_valid"] = False
                validation["errors"].append("File does not exist")
                return validation
            
            # Check file size
            file_size = os.path.getsize(file_path)
            if file_size == 0:
                validation["is_valid"] = False
                validation["errors"].append("File is empty")
                return validation
            
            if file_size > 50 * 1024 * 1024:  # 50MB
                validation["warnings"].append("File is very large (>50MB)")
            
            # Check if format is supported
            if mime_type not in self.supported_formats:
                validation["is_valid"] = False
                validation["errors"].append(f"Unsupported file type: {mime_type}")
                return validation
            
            # Try to open the file
            try:
                if mime_type == 'application/pdf':
                    with pdfplumber.open(file_path) as pdf:
                        if len(pdf.pages) == 0:
                            validation["warnings"].append("PDF has no pages")
                elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                    doc = DocxDocument(file_path)
                    if len(doc.paragraphs) == 0:
                        validation["warnings"].append("DOCX has no content")
            except Exception as e:
                validation["warnings"].append(f"Could not fully validate file: {str(e)}")
            
            return validation
            
        except Exception as e:
            logger.error(f"Error validating document: {e}")
            validation["is_valid"] = False
            validation["errors"].append(f"Validation error: {str(e)}")
            return validation
